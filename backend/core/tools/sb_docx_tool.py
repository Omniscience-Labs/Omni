"""DOCX tool – structured content blocks → .docx file.

Backend assembles a standardized JSON from LLM-provided content + overrides,
then renders it into a Word document.  Optionally uses a .docx template
(e.g. company letterhead) as the base document.
"""

from __future__ import annotations

import base64
import io
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from core.agentpress.thread_manager import ThreadManager
from core.agentpress.tool import ToolResult, openapi_schema, tool_metadata
from core.sandbox.tool_base import SandboxToolsBase
from core.utils.logger import logger

# ── Constants & defaults ─────────────────────────────────────────────────────

_PAGE_SIZES = {"LETTER": (8.5, 11.0), "A4": (8.27, 11.69)}
_HEADING_SIZES = {1: 16, 2: 14, 3: 12}

_BLOCK_DEFAULTS: Dict[str, Any] = {
    "alignment": "left",
    "bold": False,
    "italic": False,
    "underline": False,
    "color": "#000000",
    "spacing": {"before": 0, "after": 0},
}

_ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

_BULLET_CHAR = "\u2022"
_LIST_INDENT = Inches(0.5)
_HANGING_INDENT = Inches(-0.25)

# ── Helpers ──────────────────────────────────────────────────────────────────


def _validate_and_normalize_path(workspace: str, path: str) -> Tuple[str, str]:
    if not path or not isinstance(path, str):
        raise ValueError("Path must be a non-empty string")
    path = path.replace("\\", "/").strip()
    if path.startswith("/workspace/"):
        path = path[len("/workspace/"):].lstrip("/")
    normalized: List[str] = []
    for part in path.split("/"):
        if part == "..":
            if normalized:
                normalized.pop()
        elif part and part != ".":
            normalized.append(part)
    cleaned = "/".join(normalized)
    if ".." in cleaned or "\\" in cleaned:
        raise ValueError("Path cannot contain '..' or backslashes")
    return cleaned, f"{workspace}/{cleaned}"


def _sanitize_filename(name: str) -> str:
    base = "document"
    if name:
        safe = "".join(c for c in name if c.isalnum() or c in "-_.").strip()
        if safe and not safe.startswith("."):
            base = safe
    if not base.lower().endswith(".docx"):
        base = base.rstrip(".") + ".docx"
    return base


def _inches_value(value: Any, *, ceiling: float = 5.0) -> Optional[float]:
    if value is None:
        return None
    try:
        v = float(value)
    except (TypeError, ValueError):
        return None
    return v if 0 <= v <= ceiling else None


def _has_style(doc: Document, name: str) -> bool:
    try:
        doc.styles[name]
        return True
    except KeyError:
        return False


def _alignment_enum(value: str) -> Optional[WD_ALIGN_PARAGRAPH]:
    if not value:
        return None
    return _ALIGNMENT_MAP.get(str(value).strip().lower())


def _paragraph_has_text(para_element: Any) -> bool:
    """Return True if the paragraph XML element contains any non-whitespace text."""
    for el in para_element.iter():
        local = el.tag.split("}")[-1] if "}" in el.tag else el.tag
        if local == "t":
            if (el.text or "").strip() or (el.tail or "").strip():
                return True
    return False


def _strip_leading_empty_body_paragraphs(doc: Document) -> None:
    """Remove leading empty paragraphs from the document body.

    Many letterhead templates include empty paragraphs at the start of the body
    to create spacing. When we append our content, those produce an unwanted
    gap. This strips only leading empty paragraphs; leaves section properties
    and the first non-empty paragraph or table unchanged.
    """
    body = doc.element.body
    to_remove = []
    for child in list(body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "sectPr":
            break
        if tag == "p":
            if _paragraph_has_text(child):
                break
            to_remove.append(child)
        elif tag == "tbl":
            break
    for el in to_remove:
        body.remove(el)


# ── Document assembly ────────────────────────────────────────────────────────


def _build_item(item: Any, font: str, size: float, color: str) -> Dict[str, Any]:
    """Normalize a list item or table cell into a full object."""
    if isinstance(item, str):
        obj: Dict[str, Any] = {"text": item}
    elif isinstance(item, dict):
        obj = dict(item)
    else:
        obj = {"text": str(item) if item is not None else ""}
    obj.setdefault("text", "")
    obj.setdefault("font", font)
    obj.setdefault("size", size)
    obj.setdefault("color", color)
    obj.setdefault("bold", False)
    obj.setdefault("italic", False)
    obj.setdefault("underline", False)
    obj.setdefault("alignment", "left")
    return obj


def _build_block(block: Dict[str, Any], font: str, size: float, color: str) -> Dict[str, Any]:
    """Fill every missing field with defaults.  Returns a complete block."""
    b: Dict[str, Any] = dict(block)
    btype = b.get("type", "paragraph")
    b["type"] = btype

    if btype == "page_break":
        return b

    # Common defaults
    b.setdefault("font", font)
    b.setdefault("color", color)
    b.setdefault("alignment", _BLOCK_DEFAULTS["alignment"])
    b.setdefault("italic", _BLOCK_DEFAULTS["italic"])
    b.setdefault("underline", _BLOCK_DEFAULTS["underline"])

    if "spacing" not in b or not isinstance(b.get("spacing"), dict):
        b["spacing"] = dict(_BLOCK_DEFAULTS["spacing"])
    else:
        b["spacing"] = dict(b["spacing"])
        b["spacing"].setdefault("before", 0)
        b["spacing"].setdefault("after", 0)

    # Type-specific defaults
    if btype == "heading":
        level = b.get("level", 1)
        try:
            level = min(3, max(1, int(level)))
        except (TypeError, ValueError):
            level = 1
        b["level"] = level
        b.setdefault("text", "")
        b.setdefault("size", _HEADING_SIZES.get(level, 16))
        b.setdefault("bold", True)

    elif btype == "paragraph":
        b.setdefault("text", "")
        b.setdefault("size", size)
        b.setdefault("bold", False)

    elif btype == "list":
        b.setdefault("list_type", "bullet")
        b.setdefault("size", size)
        b.setdefault("bold", False)
        b["items"] = [_build_item(it, font, size, color) for it in (b.get("items") or [])]

    elif btype == "table":
        b.setdefault("size", size)
        b.setdefault("bold", False)
        raw_rows = b.get("rows") or []
        rows = [
            [_build_item(c, font, size, color) for c in (r if isinstance(r, list) else [])]
            for r in raw_rows
        ]
        if rows:
            max_cols = max((len(r) for r in rows), default=0)
            for r in rows:
                while len(r) < max_cols:
                    r.append(_build_item("", font, size, color))
        b["rows"] = rows

    else:
        b.setdefault("text", "")
        b.setdefault("size", size)
        b.setdefault("bold", False)

    return b


def _assemble_document(
    content: List[Dict[str, Any]],
    page_size: str,
    orientation: str,
    margin_top: float,
    margin_bottom: float,
    margin_left: float,
    margin_right: float,
    default_font: str,
    default_size: float,
    default_color: str,
    template_path: Optional[str] = None,
    use_template_page_setup: bool = True,
) -> Dict[str, Any]:
    """Build the full standardized document JSON from flat parameters."""
    doc_section: Dict[str, Any] = {
        "page": {
            "size": (page_size or "LETTER").upper(),
            "orientation": (orientation or "portrait").lower(),
            "margins": {
                "top": margin_top,
                "left": margin_left,
                "right": margin_right,
                "bottom": margin_bottom,
            },
        },
    }
    if template_path:
        doc_section["template"] = {
            "path": template_path,
            "use_template_page_setup": bool(use_template_page_setup),
        }
    return {
        "doc": doc_section,
        "content": [
            _build_block(b, default_font, default_size, default_color)
            for b in (content or []) if isinstance(b, dict)
        ],
    }


# ── Rendering ────────────────────────────────────────────────────────────────


def _apply_page_setup(doc: Document, cfg: Dict[str, Any]) -> None:
    section = doc.sections[0]
    w, h = _PAGE_SIZES.get(str(cfg.get("size", "LETTER")).upper(), _PAGE_SIZES["LETTER"])
    if str(cfg.get("orientation", "portrait")).lower() == "landscape":
        section.page_width, section.page_height = Inches(h), Inches(w)
        section.orientation = WD_ORIENT.LANDSCAPE
    else:
        section.page_width, section.page_height = Inches(w), Inches(h)
        section.orientation = WD_ORIENT.PORTRAIT
    margins = cfg.get("margins") or {}
    if isinstance(margins, dict):
        for key, attr in (("top", "top_margin"), ("left", "left_margin"),
                          ("right", "right_margin"), ("bottom", "bottom_margin")):
            v = _inches_value(margins.get(key))
            if v is not None:
                setattr(section, attr, Inches(v))


def _apply_run_formatting(run: Any, fmt: Dict[str, Any]) -> None:
    run.font.name = fmt["font"]
    run.font.size = Pt(fmt["size"])
    run.font.bold = fmt["bold"]
    run.font.italic = fmt["italic"]
    run.font.underline = fmt["underline"]
    hex_color = str(fmt["color"]).lstrip("#")
    if len(hex_color) == 6:
        run.font.color.rgb = RGBColor.from_string(hex_color)


def _render_blocks(doc: Document, blocks: List[Dict[str, Any]]) -> int:
    count = 0
    for blk in blocks:
        btype = blk["type"]

        if btype == "heading":
            style = f"Heading {blk['level']}"
            if _has_style(doc, style):
                para = doc.add_heading("", level=blk["level"])
                para.clear()
            else:
                para = doc.add_paragraph()
            run = para.add_run(blk["text"])
            _apply_run_formatting(run, blk)
            align = _alignment_enum(blk["alignment"])
            if align is not None:
                para.paragraph_format.alignment = align
            para.paragraph_format.space_before = Pt(blk["spacing"]["before"])
            para.paragraph_format.space_after = Pt(blk["spacing"]["after"])
            count += 1

        elif btype == "paragraph":
            para = doc.add_paragraph()
            run = para.add_run(blk["text"])
            _apply_run_formatting(run, blk)
            align = _alignment_enum(blk["alignment"])
            if align is not None:
                para.paragraph_format.alignment = align
            para.paragraph_format.space_before = Pt(blk["spacing"]["before"])
            para.paragraph_format.space_after = Pt(blk["spacing"]["after"])
            count += 1

        elif btype == "list":
            style_name = "List Bullet" if blk["list_type"] == "bullet" else "List Number"
            use_style = _has_style(doc, style_name)
            for idx, item in enumerate(blk["items"]):
                if use_style:
                    para = doc.add_paragraph(style=style_name)
                    run = para.add_run(item["text"])
                else:
                    para = doc.add_paragraph()
                    para.paragraph_format.left_indent = _LIST_INDENT
                    para.paragraph_format.first_line_indent = _HANGING_INDENT
                    prefix = f"{_BULLET_CHAR}\t" if blk["list_type"] == "bullet" else f"{idx + 1}.\t"
                    run = para.add_run(prefix + item["text"])
                _apply_run_formatting(run, item)
            count += 1

        elif btype == "table":
            rows = blk["rows"]
            if not rows:
                continue
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    cell_para = table.rows[ri].cells[ci].paragraphs[0]
                    cell_para.clear()
                    run = cell_para.add_run(cell["text"])
                    _apply_run_formatting(run, cell)
                    align = _alignment_enum(cell.get("alignment", "left"))
                    if align is not None:
                        cell_para.paragraph_format.alignment = align
            count += 1

        elif btype == "page_break":
            doc.add_page_break()
            count += 1

    return count


# ── Tool class ───────────────────────────────────────────────────────────────


@tool_metadata(
    display_name="Word Documents",
    description="Generate Word (.docx) documents from content blocks and formatting settings",
    icon="FileText",
    color="bg-indigo-100 dark:bg-indigo-800/50",
    weight=75,
    visible=True,
)
class SandboxDocxTool(SandboxToolsBase):

    def __init__(self, project_id: str, thread_manager: ThreadManager):
        super().__init__(project_id, thread_manager)

    async def _ensure_dir(self, name: str) -> None:
        try:
            await self.sandbox.fs.create_folder(f"{self.workspace_path}/{name}", "755")
        except Exception:
            pass

    # ── render_docx ──────────────────────────────────────────────────────

    @openapi_schema({
        "type": "function",
        "function": {
            "name": "render_docx",
            "description": (
                "Render a Word document from content blocks and save as .docx. "
                "Provide content as an array of blocks (heading, paragraph, list, table, page_break). "
                "The backend fills ALL missing formatting fields with defaults — only specify overrides "
                "for what the user explicitly requested. List items and table cells can be plain strings. "
                "Optionally provide template_path to use an existing .docx template file as the base "
                "(e.g. a company letterhead). The template's headers, footers, watermarks, logos, "
                "and styles are preserved, and content blocks are appended into it."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "content": {
                        "type": "array",
                        "description": (
                            "Array of content blocks. Each block must have 'type' and type-specific fields. "
                            "Types: heading (level 1-3, text), paragraph (text), list (list_type, items), "
                            "table (rows), page_break. Optional per-block overrides: font, size (pt), "
                            "color (hex e.g. '#FF0000'), bold, italic, underline, "
                            "alignment (left|center|right|justify), spacing ({before, after} in pt). "
                            "List items and table cells can be plain strings or {text, ...overrides}."
                        ),
                        "items": {
                            "type": "object",
                            "required": ["type"],
                            "properties": {
                                "type": {"type": "string", "enum": ["heading", "paragraph", "list", "table", "page_break"]},
                                "text": {"type": "string"},
                                "level": {"type": "integer", "minimum": 1, "maximum": 3},
                                "alignment": {"type": "string", "enum": ["left", "center", "right", "justify"]},
                                "font": {"type": "string"},
                                "size": {"type": "number"},
                                "color": {"type": "string"},
                                "bold": {"type": "boolean"},
                                "italic": {"type": "boolean"},
                                "underline": {"type": "boolean"},
                                "spacing": {
                                    "type": "object",
                                    "properties": {
                                        "before": {"type": "number"},
                                        "after": {"type": "number"},
                                    },
                                },
                                "list_type": {"type": "string", "enum": ["bullet", "number"]},
                                "items": {"type": "array", "description": "List items: strings or {text, font?, size?, bold?, italic?, underline?, color?}"},
                                "rows": {"type": "array", "description": "Table rows: array of arrays. Each cell: string or {text, font?, size?, bold?, italic?, underline?, color?, alignment?}"},
                            },
                        },
                    },
                    "output_filename": {"type": "string", "description": "Output filename (e.g. 'report.docx').", "default": "document.docx"},
                    "template_path": {
                        "type": "string",
                        "description": (
                            "Path to a .docx template file in the sandbox "
                            "(e.g. '/workspace/templates/letterhead.docx'). When provided, "
                            "the template is used as the base document — its headers, footers, "
                            "watermarks, logos, and styles are preserved. Content blocks are "
                            "rendered into the template's body. Omit or set to null to create "
                            "a blank document (default behaviour)."
                        ),
                    },
                    "use_template_page_setup": {
                        "type": "boolean",
                        "description": (
                            "When true (default), page size/orientation/margins are inherited "
                            "from the template. Set to false to override with the explicit "
                            "page_size/orientation/margin_* parameters."
                        ),
                        "default": True,
                    },
                    "page_size": {"type": "string", "enum": ["LETTER", "A4"], "description": "Page size.", "default": "LETTER"},
                    "orientation": {"type": "string", "enum": ["portrait", "landscape"], "description": "Page orientation.", "default": "portrait"},
                    "margin_top": {"type": "number", "description": "Top margin in inches.", "default": 1.0},
                    "margin_bottom": {"type": "number", "description": "Bottom margin in inches.", "default": 1.0},
                    "margin_left": {"type": "number", "description": "Left margin in inches.", "default": 1.0},
                    "margin_right": {"type": "number", "description": "Right margin in inches.", "default": 1.0},
                    "default_font": {"type": "string", "description": "Default font for all blocks.", "default": "Calibri"},
                    "default_font_size": {"type": "number", "description": "Default font size in points.", "default": 11},
                    "default_font_color": {"type": "string", "description": "Default font color as hex.", "default": "#000000"},
                },
                "required": ["content"],
            },
        },
    })
    async def render_docx(
        self,
        content: List[Dict[str, Any]],
        output_filename: str = "document.docx",
        template_path: Optional[str] = None,
        use_template_page_setup: bool = True,
        page_size: str = "LETTER",
        orientation: str = "portrait",
        margin_top: float = 1.0,
        margin_bottom: float = 1.0,
        margin_left: float = 1.0,
        margin_right: float = 1.0,
        default_font: str = "Calibri",
        default_font_size: float = 11.0,
        default_font_color: str = "#000000",
    ) -> ToolResult:
        try:
            await self._ensure_sandbox()
            await self._ensure_dir("documents")

            if not content or not isinstance(content, list):
                return self.fail_response("'content' must be a non-empty array of block objects.")

            # Validate template_path
            resolved_template: Optional[str] = None
            if template_path:
                if not template_path.strip().lower().endswith(".docx"):
                    return self.fail_response("Template must be a .docx file.")
                _, resolved_template = _validate_and_normalize_path(self.workspace_path, template_path)

            # Assemble standardized document JSON
            assembled = _assemble_document(
                content=content,
                page_size=page_size,
                orientation=orientation,
                margin_top=margin_top,
                margin_bottom=margin_bottom,
                margin_left=margin_left,
                margin_right=margin_right,
                default_font=default_font,
                default_size=default_font_size,
                default_color=default_font_color,
                template_path=template_path,
                use_template_page_setup=use_template_page_setup,
            )
            page_cfg = assembled["doc"]["page"]
            tmpl_cfg = assembled["doc"].get("template")
            blocks = assembled["content"]

            for blk in blocks:
                if blk["type"] == "table" and not blk.get("rows"):
                    return self.fail_response("Table block must have at least one row.")

            # Build the .docx
            using_template = False
            if tmpl_cfg and resolved_template:
                try:
                    raw = await self.sandbox.fs.download_file(resolved_template)
                    doc = Document(io.BytesIO(raw))
                    using_template = True
                    _strip_leading_empty_body_paragraphs(doc)
                    logger.info("sb_docx_tool: loaded template %s", resolved_template)
                except Exception as e:
                    return self.fail_response(f"Failed to load template '{template_path}': {e}")
                if not tmpl_cfg.get("use_template_page_setup", True):
                    _apply_page_setup(doc, page_cfg)
            else:
                doc = Document()
                _apply_page_setup(doc, page_cfg)

            block_count = _render_blocks(doc, blocks)

            # Save
            safe_name = _sanitize_filename(output_filename)
            cleaned, full_path = _validate_and_normalize_path(self.workspace_path, f"documents/{safe_name}")
            try:
                await self.sandbox.fs.delete_file(full_path)
            except Exception:
                pass
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            await self.sandbox.fs.upload_file(buf.getvalue(), full_path)

            logger.info("sb_docx_tool: saved %s (%d blocks, template=%s)", full_path, block_count, using_template)
            result: Dict[str, Any] = {
                "relative_path": f"/workspace/{cleaned}",
                "block_count": block_count,
                "template_used": using_template,
            }
            if using_template:
                result["template_source"] = template_path
            return self.success_response(result)

        except ValueError as e:
            return self.fail_response(str(e))
        except Exception as e:
            logger.exception("sb_docx_tool render_docx failed")
            return self.fail_response(f"Failed to render document: {e}")

    # ── upload_template ──────────────────────────────────────────────────

    @openapi_schema({
        "type": "function",
        "function": {
            "name": "upload_template",
            "description": (
                "Upload a .docx file as a reusable document template "
                "(e.g. a company letterhead). Stored in /workspace/templates/ "
                "and usable via template_path in render_docx."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "Name for the template file (e.g. 'letterhead.docx'). Must end in .docx.",
                    },
                    "file_base64": {
                        "type": "string",
                        "description": "Base64-encoded content of the .docx file.",
                    },
                },
                "required": ["filename", "file_base64"],
            },
        },
    })
    async def upload_template(self, filename: str, file_base64: str) -> ToolResult:
        try:
            await self._ensure_sandbox()
            await self._ensure_dir("templates")

            if not filename or not isinstance(filename, str):
                return self.fail_response("'filename' must be a non-empty string.")
            if not filename.strip().lower().endswith(".docx"):
                return self.fail_response("Template must be a .docx file.")

            try:
                file_bytes = base64.b64decode(file_base64)
            except Exception:
                return self.fail_response("Invalid base64 data.")
            if not file_bytes:
                return self.fail_response("Template file content is empty.")

            # Validate it's a real Word file
            try:
                Document(io.BytesIO(file_bytes))
            except Exception:
                return self.fail_response("Not a valid .docx file.")

            safe = "".join(c for c in filename if c.isalnum() or c in "-_.").strip() or "template.docx"
            if not safe.lower().endswith(".docx"):
                safe = safe.rstrip(".") + ".docx"

            cleaned, full_path = _validate_and_normalize_path(self.workspace_path, f"templates/{safe}")
            try:
                await self.sandbox.fs.delete_file(full_path)
            except Exception:
                pass
            await self.sandbox.fs.upload_file(file_bytes, full_path)

            logger.info("sb_docx_tool: uploaded template %s (%d bytes)", full_path, len(file_bytes))
            return self.success_response({
                "template_path": f"/workspace/{cleaned}",
                "size_bytes": len(file_bytes),
            })

        except ValueError as e:
            return self.fail_response(str(e))
        except Exception as e:
            logger.exception("sb_docx_tool upload_template failed")
            return self.fail_response(f"Failed to upload template: {e}")

    # ── list_templates ───────────────────────────────────────────────────

    @openapi_schema({
        "type": "function",
        "function": {
            "name": "list_templates",
            "description": (
                "List available document templates (.docx) in /workspace/templates/. "
                "Returns paths usable as template_path in render_docx."
            ),
            "parameters": {"type": "object", "properties": {}},
        },
    })
    async def list_templates(self) -> ToolResult:
        try:
            await self._ensure_sandbox()
            await self._ensure_dir("templates")

            try:
                entries = await self.sandbox.fs.list_dir(f"{self.workspace_path}/templates")
            except Exception:
                return self.success_response({"templates": [], "count": 0})

            templates = []
            for entry in entries:
                name = entry.name if hasattr(entry, "name") else str(entry)
                if name.strip().lower().endswith(".docx"):
                    templates.append({
                        "filename": name,
                        "template_path": f"/workspace/templates/{name}",
                    })

            return self.success_response({"templates": templates, "count": len(templates)})

        except Exception as e:
            logger.exception("sb_docx_tool list_templates failed")
            return self.fail_response(f"Failed to list templates: {e}")
